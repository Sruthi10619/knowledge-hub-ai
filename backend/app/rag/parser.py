"""
Multi-format document parser.
Supports PDF, DOCX, TXT, MD, CSV.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Any, Dict, List

import pdfplumber
import PyPDF2
from docx import Document as DocxDocument

from app.core.exceptions import DocumentProcessingError


class DocumentParser:
    @classmethod
    def parse(cls, filename: str, content: bytes) -> List[Dict[str, Any]]:
        """
        Parses a document based on its extension.
        Returns a list of dictionaries containing:
        - "text": content of the section/chunk
        - "metadata": dict containing page, line, source, etc.
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        try:
            if ext == "pdf":
                return cls._parse_pdf(filename, content)
            elif ext == "docx":
                return cls._parse_docx(filename, content)
            elif ext == "csv":
                return cls._parse_csv(filename, content)
            elif ext in ("txt", "md"):
                return cls._parse_text(filename, content)
            else:
                raise DocumentProcessingError(f"Unsupported file format: {ext}")
        except Exception as e:
            if isinstance(e, DocumentProcessingError):
                raise e
            raise DocumentProcessingError(f"Error parsing {filename}: {str(e)}")

    @classmethod
    def _parse_pdf(cls, filename: str, content: bytes) -> List[Dict[str, Any]]:
        results = []
        # Try pdfplumber first (higher quality extraction)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for idx, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        results.append({
                            "text": text,
                            "metadata": {
                                "source": filename,
                                "page": idx + 1,
                                "total_pages": len(pdf.pages)
                            }
                        })
            if results:
                return results
        except Exception:
            # Fallback to PyPDF2
            pass

        # Fallback PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        total_pages = len(pdf_reader.pages)
        for idx in range(total_pages):
            page = pdf_reader.pages[idx]
            text = page.extract_text()
            if text and text.strip():
                results.append({
                    "text": text,
                    "metadata": {
                        "source": filename,
                        "page": idx + 1,
                        "total_pages": total_pages
                    }
                })
        return results

    @classmethod
    def _parse_docx(cls, filename: str, content: bytes) -> List[Dict[str, Any]]:
        results = []
        doc = DocxDocument(io.BytesIO(content))
        
        # Group paragraphs by sections or just accumulate
        full_text = []
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                full_text.append((idx + 1, text))
                
        # To avoid creating too many small chunks or single large one, we return sections
        # For simplicity in parsing, we export paragraphs with line metadata
        # which will be grouped by our Recursive Chunker.
        combined_text = "\n".join([t[1] for t in full_text])
        if combined_text.strip():
            results.append({
                "text": combined_text,
                "metadata": {
                    "source": filename,
                    "total_paragraphs": len(doc.paragraphs)
                }
            })
        return results

    @classmethod
    def _parse_csv(cls, filename: str, content: bytes) -> List[Dict[str, Any]]:
        results = []
        text_stream = io.StringIO(content.decode("utf-8", errors="ignore"))
        reader = csv.reader(text_stream)
        
        headers = next(reader, None)
        if not headers:
            return results

        # Parse row by row, building structured text representations
        for idx, row in enumerate(reader):
            row_items = []
            for header, val in zip(headers, row):
                row_items.append(f"{header}: {val}")
            
            row_text = ", ".join(row_items)
            results.append({
                "text": row_text,
                "metadata": {
                    "source": filename,
                    "row": idx + 1,
                    "headers": headers
                }
            })
        return results

    @classmethod
    def _parse_text(cls, filename: str, content: bytes) -> List[Dict[str, Any]]:
        text = content.decode("utf-8", errors="ignore")
        if not text.strip():
            return []
            
        return [{
            "text": text,
            "metadata": {
                "source": filename
            }
        }]
